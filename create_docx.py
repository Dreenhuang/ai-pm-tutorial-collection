"""
Create formatted .docx document for AI Programming PM Tutorial Collection
Format requirements:
- Font: SimSun (宋体)
- Font size: 4号 (14pt)
- Line spacing: 22pt (exactly)
- Page numbers: continuous from 1 at bottom
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

# Configuration
BASE_DIR = r"c:\Users\Administrator\Documents\xuexi\xuexi-trae2"
TUTORIALS_DIR = os.path.join(BASE_DIR, "tutorials")
OUTPUT_FILE = os.path.join(BASE_DIR, "教程全集_完整版.docx")

# Formatting constants
FONT_NAME = '宋体'
FONT_SIZE = Pt(14)  # 4号 = 14pt
LINE_SPACING = Pt(22)  # 22pt fixed

HEADING1_SIZE = Pt(22)
HEADING2_SIZE = Pt(18)
HEADING3_SIZE = Pt(16)
HEADING4_SIZE = Pt(14)

def create_document():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Set East Asian font
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {
        qn('w:eastAsia'): FONT_NAME,
        qn('w:ascii'): FONT_NAME,
        qn('w:hAnsi'): FONT_NAME
    })
    rPr.append(rFonts)
    
    # Set paragraph spacing
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
        
        # Add page numbers in footer
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number field
        run = footer_para.add_run()
        fldChar1 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
        run._element.append(fldChar1)
        
        run2 = footer_para.add_run()
        instrText = run2._element.makeelement(qn('w:instrText'), {})
        instrText.text = ' PAGE '
        run2._element.append(instrText)
        
        run3 = footer_para.add_run()
        fldChar2 = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
        run3._element.append(fldChar2)
        
        run.font.size = Pt(10.5)  # 5号
        run.font.name = FONT_NAME
        run2.font.size = Pt(10.5)
        run2.font.name = FONT_NAME
        run3.font.size = Pt(10.5)
        run3.font.name = FONT_NAME
    
    return doc

def add_heading_styled(doc, text, level):
    """Add heading with custom formatting"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER
    
    if level == 1:
        size = HEADING1_SIZE
        bold = True
        para.paragraph_format.space_before = Pt(24)
        para.paragraph_format.space_after = Pt(12)
    elif level == 2:
        size = HEADING2_SIZE
        bold = True
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(9)
    elif level == 3:
        size = HEADING3_SIZE
        bold = True
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
    else:
        size = HEADING4_SIZE
        bold = True
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(4)
    
    run = para.add_run(text)
    run.font.size = size
    run.bold = bold
    run.font.name = FONT_NAME
    
    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {
        qn('w:eastAsia'): FONT_NAME,
        qn('w:ascii'): 'Arial',
        qn('w:hAnsi'): 'Arial'
    })
    rPr.append(rFonts)
    
    return para

def add_text(doc, text, bold=False, indent=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Add formatted text paragraph"""
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.line_spacing = LINE_SPACING
    
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)  # 2字符缩进
    
    run = para.add_run(text)
    run.font.size = FONT_SIZE
    run.bold = bold
    run.font.name = FONT_NAME
    
    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {
        qn('w:eastAsia'): FONT_NAME,
        qn('w:ascii'): 'Arial',
        qn('w:hAnsi'): 'Arial'
    })
    rPr.append(rFonts)
    
    return para

def add_table_from_md(doc, rows_data):
    """Add table from markdown data"""
    if not rows_data or len(rows_data) < 2:
        return
    
    headers = rows_data[0]
    data_rows = rows_data[1:]
    
    table = doc.add_table(rows=len(data_rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Add headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = FONT_NAME
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add data
    for row_idx, row_data in enumerate(data_rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(table.rows[row_idx+1].cells):
                cell = table.rows[row_idx+1].cells[col_idx]
                cell.text = str(cell_text)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(11)
                        run.font.name = FONT_NAME
    
    return table

def parse_tutorial_file(filepath):
    """Parse a markdown tutorial file and return structured content"""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    sections = []
    current_section = {'title': '', 'content': []}
    
    for line in content.split('\n'):
        line = line.rstrip()
        
        if line.startswith('# '):
            if current_section['title']:
                sections.append(current_section)
            current_section = {'title': line[2:].strip(), 'content': [], 'level': 1}
        elif line.startswith('## '):
            if current_section['title']:
                sections.append(current_section)
            current_section = {'title': line[3:].strip(), 'content': [], 'level': 2}
        elif line.startswith('### '):
            if current_section['title']:
                sections.append(current_section)
            current_section = {'title': line[4:].strip(), 'content': [], 'level': 3}
        elif line.startswith('#### '):
            if current_section['title']:
                sections.append(current_section)
            current_section = {'title': line[5:].strip(), 'content': [], 'level': 4}
        elif line.strip():
            current_section['content'].append(line)
    
    if current_section['title']:
        sections.append(current_section)
    
    return sections

def clean_text(text):
    """Remove markdown formatting from text"""
    # Remove markdown links but keep text
    import re
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # Remove bold/italic markers
    text = text.replace('**', '').replace('*', '')
    # Remove blockquote markers
    text = text.replace('> ', '')
    # Remove horizontal rules
    if text.strip() == '---':
        return ''
    return text

def main():
    print("Creating .docx document...")
    doc = create_document()
    
    # Add cover page
    add_text(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    add_heading_styled(doc, 'AI编程时代产品经理入门教程', 1)
    add_heading_styled(doc, '完整版', 1)
    
    add_text(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, '目标读者：零技术基础或刚入门的产品经理', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, '学习周期：8-12周（每周10-15小时）', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, '总教程数：110个子项', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, '创建日期：2026-04-26', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, '最后更新：2026-04-27', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Page break after cover
    doc.add_page_break()
    
    # Process each module
    modules = [
        ('模块一', '互联网产品基础认知', os.path.join(TUTORIALS_DIR, 'module01')),
        ('模块二', '前端开发核心概念', os.path.join(TUTORIALS_DIR, 'module02')),
        ('模块三', '后端开发核心概念', os.path.join(TUTORIALS_DIR, 'module03')),
        ('模块四', '产品运营与数据驱动', os.path.join(TUTORIALS_DIR, 'module04')),
        ('模块五', 'AI编程与AI产品核心概念', os.path.join(TUTORIALS_DIR, 'module05')),
        ('模块六', 'AI编程实战工作流', os.path.join(TUTORIALS_DIR, 'module06')),
        ('模块七', '进阶专题与行业趋势', os.path.join(TUTORIALS_DIR, 'module07')),
    ]
    
    file_count = 0
    
    for module_num, module_name, module_dir in modules:
        if not os.path.exists(module_dir):
            print(f"Warning: {module_dir} does not exist, skipping...")
            continue
        
        # Add module heading
        add_heading_styled(doc, f'{module_num}：{module_name}', 1)
        add_text(doc, '')
        
        # Get all markdown files
        md_files = sorted([f for f in os.listdir(module_dir) if f.endswith('.md')])
        
        for md_file in md_files:
            filepath = os.path.join(module_dir, md_file)
            file_count += 1
            
            # Count total files
            total_files = 0
            for d in os.listdir(TUTORIALS_DIR):
                full_d = os.path.join(TUTORIALS_DIR, d)
                if os.path.isdir(full_d):
                    total_files += len([f for f in os.listdir(full_d) if f.endswith('.md')])
            
            print(f"Processing: {md_file} ({file_count}/{total_files})")
            
            sections = parse_tutorial_file(filepath)
            
            for section in sections:
                # Skip metadata block
                if '所属模块' in section['title'] or '---' in section['title']:
                    continue
                
                # Add section heading
                level = min(section.get('level', 3), 3)
                add_heading_styled(doc, section['title'], level)
                
                # Add content
                for line in section['content']:
                    cleaned = clean_text(line)
                    if not cleaned:
                        continue
                    
                    # Check if it's a list item
                    if cleaned.startswith('- ') or cleaned.startswith('* '):
                        add_text(doc, cleaned[2:], indent=True)
                    elif cleaned.startswith('|') and '---' not in cleaned:
                        # Table row - for simplicity, add as text
                        cells = [c.strip() for c in cleaned.split('|') if c.strip()]
                        if cells:
                            add_text(doc, ' | '.join(cells), bold=cleaned.startswith('| **'))
                    elif cleaned.startswith('```'):
                        # Code block
                        add_text(doc, cleaned, bold=True)
                    else:
                        add_text(doc, cleaned, indent=True)
                
                add_text(doc, '')
            
            # Add page break between tutorials
            doc.add_page_break()
        
        # Page break between modules
        doc.add_page_break()
    
    print(f"\nSaving document to: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)
    print(f"✅ Document created successfully with {file_count} tutorials!")
    print(f"📄 Output file: {OUTPUT_FILE}")

if __name__ == '__main__':
    main()
